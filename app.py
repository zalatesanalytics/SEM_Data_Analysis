
from pathlib import Path
import re
import matplotlib.pyplot as plt
import networkx as nx
import pandas as pd
import seaborn as sns
import streamlit as st
try:
    from scipy import stats
except Exception:
    stats = None
from docx import Document
from docx.shared import Inches

from sem_engine import (
    load_data, normalize_columns, coerce_analysis_columns, variable_mapping,
    descriptive_stats, reliability_table, construct_scores, correlation_matrix,
    path_coefficients, mediation_analysis, semopy_model, write_excel,
)
from sem_config import LATENT_VARIABLES

st.set_page_config(page_title="EMERGE+ SEM Analyzer", layout="wide")
st.title("EMERGE+ Newcomer SEM Analyzer")
st.caption(
    "Upload a CSV or Excel file, clean variables, estimate SEM pathways, "
    "visualize results, generate descriptive analysis, and produce a Word report."
)

OUTPUT_ROOT = Path("outputs")
for sub in ["tables", "figures", "processed_data", "model_outputs"]:
    (OUTPUT_ROOT / sub).mkdir(parents=True, exist_ok=True)

st.markdown("""
<style>
.main .block-container {padding-top: 1.2rem; padding-bottom: 2rem;}
.section-card {
    background: #ffffff;
    border: 1px solid #e8eef7;
    border-radius: 16px;
    padding: 1rem 1.2rem;
    box-shadow: 0 4px 14px rgba(28, 39, 60, 0.06);
    margin-bottom: 0.8rem;
}
.metric-card {
    background: linear-gradient(135deg, #f7fbff 0%, #eef5ff 100%);
    border: 1px solid #dce9fa;
    border-radius: 14px;
    padding: 0.9rem;
    text-align: center;
}
.metric-card h3 {margin: 0; font-size: 0.9rem; color: #4a5568;}
.metric-card p {margin: 0.2rem 0 0 0; font-size: 1.6rem; font-weight: 700; color: #1f3b63;}
.insight-box {
    background: #f8fafc;
    border-left: 5px solid #2f80ed;
    border-radius: 10px;
    padding: 0.8rem 1rem;
    margin: 0.5rem 0 1rem 0;
}
.small-caption {font-size: 0.85rem; color: #5f6b7a;}
.hero-card {
    background: linear-gradient(135deg, #0f766e 0%, #134e4a 100%);
    color: white;
    border-radius: 20px;
    padding: 1.2rem 1.4rem;
    box-shadow: 0 8px 22px rgba(15, 118, 110, 0.22);
    margin-bottom: 1rem;
}
.hero-card h2 {margin: 0; color: white;}
.hero-card p {margin: 0.35rem 0 0 0; color: #e6fffb;}
.stat-strip {
    background: #ffffff;
    border: 1px solid #e5eef5;
    border-radius: 14px;
    padding: 0.85rem 1rem;
    box-shadow: 0 3px 12px rgba(15, 23, 42, 0.06);
    margin-bottom: 0.75rem;
}
.stat-strip b {color: #0f766e;}
.dashboard-divider {height: 1px; background: #e5e7eb; margin: 1rem 0;}
</style>
""", unsafe_allow_html=True)


# -------------------------------------------------------------------
# General helper functions
# -------------------------------------------------------------------
def safe_filename(name: str) -> str:
    """Convert any variable name into a safe file name."""
    return re.sub(r"[^A-Za-z0-9_]+", "_", str(name))[:80]


def pretty_label(name: str) -> str:
    """Create a reader-friendly variable label for dashboard titles and tables."""
    label_map = {
        "gender": "Gender",
        "sex": "Gender",
        "gender_group": "Gender",
        "year_arrival": "Year of Arrival",
        "arrival_year": "Year of Arrival",
        "year_of_arrival": "Year of Arrival",
        "country_origin": "Country of Origin",
        "country_of_origin": "Country of Origin",
        "immigration_status": "Immigration Status",
        "immigration_status_us": "Immigration Status",
        "marital_status": "Marital Status",
        "highest_education": "Highest Education",
        "highest_education_level": "Highest Education",
        "employment_status": "Employment Status",
        "employment_status_us": "Employment Status",
        "annual_income_numeric": "Annual Income",
        "annual_income": "Annual Income",
    }
    key = str(name).strip().lower()
    return label_map.get(key, str(name).replace("_", " ").title())


def find_column_by_aliases(df, aliases):
    """Find a column using common aliases after normalization."""
    lower_map = {str(col).lower().strip(): col for col in df.columns}
    for alias in aliases:
        if alias.lower().strip() in lower_map:
            return lower_map[alias.lower().strip()]
    for col in df.columns:
        compact = str(col).lower().replace("_", "").replace(" ", "")
        for alias in aliases:
            alias_compact = alias.lower().replace("_", "").replace(" ", "")
            if compact == alias_compact:
                return col
    return None


def variable_distribution_table(df, col):
    """Create a clean distribution table without repeating a Variable column."""
    if col is None or col not in df.columns:
        return pd.DataFrame()

    display_col = pretty_label(col)
    counts = df[col].fillna("Missing").astype(str).value_counts(dropna=False)
    total = counts.sum()

    return pd.DataFrame({
        display_col: counts.index,
        "Frequency": counts.values.astype(int),
        "Percentage": ((counts.values / total) * 100).round(2) if total else 0,
    })


def build_key_distribution_tables(df, gender_col=None):
    """Build clean distribution tables for key participant profile variables."""
    key_specs = {
        "Gender": ["gender", "sex", "respondent_gender", "participant_gender", "gender_identity"],
        "Year of Arrival": ["year_arrival", "arrival_year", "year_of_arrival", "year_arrived", "arrival"],
        "Country of Origin": ["country_origin", "country_of_origin", "origin_country", "country"],
        "Immigration Status": ["immigration_status", "immigration_stat_us", "immigration_status_us", "immigration_stat", "status"],
        "Marital Status": ["marital_status", "marital"],
        "Highest Education": ["highest_education", "highest_educatio_n", "highest_education_level", "education", "education_level"],
        "Employment Status": ["employment_status", "employment_stat_us", "employment_status_us", "employment_stat", "employment"],
    }

    tables = {}
    used_cols = set()
    for title, aliases in key_specs.items():
        col = gender_col if title == "Gender" and gender_col in df.columns else find_column_by_aliases(df, aliases)
        if col is not None and col not in used_cols:
            tables[title] = {"column": col, "table": variable_distribution_table(df, col)}
            used_cols.add(col)
    return tables


def add_docx_table(doc, df: pd.DataFrame):
    """Add a pandas DataFrame to a Word document."""
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return

    display = df.copy().fillna("")
    table = doc.add_table(rows=1, cols=len(display.columns))
    table.style = "Table Grid"

    for i, col in enumerate(display.columns):
        table.rows[0].cells[i].text = str(col)

    for _, row in display.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)[:120]


def detect_numeric_variables(df):
    """Detect quantitative variables."""
    return df.select_dtypes(include="number").columns.tolist()


def detect_likert_variables(df):
    """
    Detect Likert-scale variables coded 1 to 5.
    These are used to calculate percentage agreeing or strongly agreeing.
    """
    likert_cols = []
    for col in df.columns:
        values = df[col].dropna().unique()
        numeric_values = pd.to_numeric(pd.Series(values), errors="coerce").dropna()
        if len(numeric_values) > 0 and numeric_values.between(1, 5).all():
            unique_count = numeric_values.nunique()
            if 2 <= unique_count <= 5:
                likert_cols.append(col)
    return likert_cols


def detect_yes_no_variables(df):
    """Detect Yes/No variables."""
    yes_no_cols = []
    yes_no_values = {"yes", "no", "1", "0", "true", "false"}

    for col in df.columns:
        values = df[col].dropna().astype(str).str.lower().str.strip().unique()
        if len(values) > 0 and set(values).issubset(yes_no_values):
            yes_no_cols.append(col)

    return yes_no_cols


def detect_categorical_variables(df, numeric_cols=None, max_unique=30):
    """
    Detect categorical variables for frequency distribution.
    Numeric columns with very few unique values are also allowed as categorical.
    """
    numeric_cols = numeric_cols or []
    categorical_cols = []

    for col in df.columns:
        unique_count = df[col].nunique(dropna=True)

        if unique_count == 0:
            continue

        if col not in numeric_cols:
            if unique_count <= max_unique:
                categorical_cols.append(col)
        else:
            if unique_count <= 10:
                categorical_cols.append(col)

    return categorical_cols


def find_gender_column(df):
    """
    Find the most likely gender column.
    Priority is given to columns named gender, sex, respondent_gender, or participant_gender.
    """
    candidate_names = [
        "gender", "sex", "respondent_gender", "participant_gender",
        "gender_identity", "respondent_sex"
    ]

    lower_map = {col.lower().strip(): col for col in df.columns}

    for name in candidate_names:
        if name in lower_map:
            return lower_map[name]

    for col in df.columns:
        if "gender" in col.lower() or col.lower().strip() == "sex":
            return col

    return None


def standardize_gender_for_men_women(df, gender_col):
    """
    Keep gender disaggregation only for Men and Women.
    Any third option or other response is excluded from gender-disaggregated analysis.
    """
    if gender_col is None or gender_col not in df.columns:
        return df.copy(), None

    df_gender = df.copy()

    def recode_gender(value):
        if pd.isna(value):
            return None

        text = str(value).strip().lower()

        if text in ["male", "man", "men", "m", "1"]:
            return "Men"
        if text in ["female", "woman", "women", "f", "2"]:
            return "Women"

        return None

    df_gender["gender_group"] = df_gender[gender_col].apply(recode_gender)
    df_gender = df_gender[df_gender["gender_group"].isin(["Men", "Women"])].copy()

    return df_gender, "gender_group"




def find_column_by_candidates(df, candidates):
    """Find a column using expected names, allowing minor case/spacing differences."""
    if df is None or df.empty:
        return None

    normalized_map = {
        re.sub(r"[^a-z0-9]+", "_", str(col).strip().lower()).strip("_"): col
        for col in df.columns
    }

    for candidate in candidates:
        key = re.sub(r"[^a-z0-9]+", "_", str(candidate).strip().lower()).strip("_")
        if key in normalized_map:
            return normalized_map[key]

    for key, original in normalized_map.items():
        for candidate in candidates:
            cand_key = re.sub(r"[^a-z0-9]+", "_", str(candidate).strip().lower()).strip("_")
            if cand_key and (cand_key in key or key in cand_key):
                return original

    return None



def income_band_to_numeric(value):
    """
    Convert annual income bands such as 25_49k, 50_74k, 75_99k, 100k_plus
    into approximate numeric annual income midpoints.

    This keeps the original annual_income column unchanged and creates
    an analysis-ready numeric version for descriptives, gender comparison,
    cost-of-stalling, charts, and SEM/correlation-ready summaries.
    """
    if pd.isna(value):
        return pd.NA

    text = str(value).strip().lower()
    text = text.replace("$", "").replace(",", "").replace(" ", "")
    text = text.replace("-", "_").replace("to", "_")

    mapping = {
        "0_24k": 12000,
        "under_25k": 12000,
        "less_than_25k": 12000,
        "<25k": 12000,
        "25_49k": 37000,
        "25k_49k": 37000,
        "25_50k": 37500,
        "50_74k": 62000,
        "50k_74k": 62000,
        "50_75k": 62500,
        "75_99k": 87000,
        "75k_99k": 87000,
        "75_100k": 87500,
        "100k_plus": 110000,
        "100k+": 110000,
        "100_plus": 110000,
        "over_100k": 110000,
        ">100k": 110000,
        "prefer_not": pd.NA,
        "prefer_not_to_say": pd.NA,
        "dont_know": pd.NA,
        "missing": pd.NA,
        "nan": pd.NA,
    }

    if text in mapping:
        return mapping[text]

    # If already numeric, keep it.
    numeric = pd.to_numeric(text, errors="coerce")
    if pd.notna(numeric):
        # Treat small values like 50 or 75 as thousands if needed.
        return numeric * 1000 if numeric < 1000 else numeric

    # Flexible pattern: 25_49k, 25k_49k, etc.
    nums = re.findall(r"\d+", text)
    if len(nums) >= 2:
        low = float(nums[0])
        high = float(nums[1])
        midpoint = (low + high) / 2
        return midpoint * 1000 if "k" in text or midpoint < 1000 else midpoint

    if len(nums) == 1 and ("plus" in text or "+" in text or "over" in text or ">" in text):
        val = float(nums[0])
        return val * 1000 if "k" in text or val < 1000 else val

    return pd.NA


def add_income_numeric_columns(df):
    """
    Add numeric income columns when income is stored as categories.
    This fixes cases where annual_income exists but does not appear in quantitative analysis
    because values are text bands such as 25_49k or 100k_plus.
    """
    df = df.copy()

    income_candidates = [
        "annual_income",
        "income",
        "current_income",
        "household_income",
        "personal_income",
        "employment_income"
    ]

    for col in income_candidates:
        if col in df.columns:
            numeric_col = f"{col}_numeric"
            converted = df[col].apply(income_band_to_numeric)
            converted = pd.to_numeric(converted, errors="coerce")

            # Only add when conversion produces usable values.
            if converted.notna().sum() > 0:
                df[numeric_col] = converted

    return df


def key_dataset_statistics(df):
    """Create high-level data readiness statistics for the home page."""
    total_cells = int(df.shape[0] * df.shape[1]) if df is not None else 0
    missing_cells = int(df.isna().sum().sum()) if df is not None else 0
    missing_pct = round((missing_cells / total_cells) * 100, 2) if total_cells else 0

    return pd.DataFrame([
        {"Statistic": "Number of observations / rows", "Value": f"{df.shape[0]:,}"},
        {"Statistic": "Number of columns / variables", "Value": f"{df.shape[1]:,}"},
        {"Statistic": "Duplicated rows", "Value": f"{int(df.duplicated().sum()):,}"},
        {"Statistic": "Total missing values", "Value": f"{missing_cells:,}"},
        {"Statistic": "Missing values as % of all cells", "Value": f"{missing_pct}%"},
    ])


def missing_values_summary(df):
    """Summarize missing values by variable."""
    if df is None or df.empty:
        return pd.DataFrame()

    miss = df.isna().sum().reset_index()
    miss.columns = ["Variable", "Missing values"]
    miss["Missing percentage"] = (miss["Missing values"] / len(df) * 100).round(2) if len(df) else 0
    miss = miss.sort_values(["Missing values", "Variable"], ascending=[False, True])
    return miss


def distribution_for_column(df, col, label=None):
    """Frequency and percentage table for a selected categorical column without a repeated Variable column."""
    if col is None or col not in df.columns:
        return pd.DataFrame()

    display_col = pretty_label(label or col)
    series = df[col].fillna("Missing").astype(str).str.strip()
    counts = series.value_counts(dropna=False)
    total = counts.sum()

    return pd.DataFrame({
        display_col: counts.index,
        "Frequency": counts.values.astype(int),
        "Percentage": ((counts.values / total) * 100).round(2) if total else 0,
    })


def priority_distribution_tables(df, gender_col=None):
    """Generate priority demographic and background distributions requested for the dashboard."""
    requested = {
        "Gender": ["gender", "sex", "respondent_gender", "participant_gender", "gender_identity"],
        "Year of arrival": ["year_arrival", "year_of_arrival", "arrival_year", "year_arrived", "arrival_date", "year_arrive"],
        "Country of origin": ["country_origin", "country_of_origin", "origin_country", "birth_country", "country_birth"],
        "Immigration status": ["immigration_status", "immigration_stat_us", "immigration_stat", "immigration_category", "legal_status"],
        "Marital status": ["marital_status", "marital", "relationship_status"],
        "Highest education": ["highest_education", "highest_educatio_n", "education", "education_level", "highest_level_education"],
        "Employment status": ["employment_status", "employment_stat_us", "employment_stat", "work_status", "current_employment_status"],
    }

    tables = {}
    detected = {}

    for label, candidates in requested.items():
        col = gender_col if label == "Gender" and gender_col is not None else find_column_by_candidates(df, candidates)
        detected[label] = col
        tables[label] = distribution_for_column(df, col, label=label)

    return tables, detected


def combine_priority_distributions(priority_tables):
    """Combine requested distribution tables into one downloadable table.

    The dashboard displays each distribution separately without a repeated variable column.
    This combined table adds a Distribution column only for Excel/Word export clarity.
    """
    frames = []
    for label, table in priority_tables.items():
        if table is None or table.empty:
            continue
        temp = table.copy()
        first_col = [c for c in temp.columns if c not in ["Frequency", "Percentage"]][0]
        temp = temp.rename(columns={first_col: "Category"})
        temp.insert(0, "Distribution", label)
        frames.append(temp)
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()


def save_priority_distribution_charts(priority_tables, output_dir):
    """Save compact bar charts for priority demographic/background distributions."""
    saved_paths = []

    for label, table in priority_tables.items():
        if table is None or table.empty:
            continue

        category_col = [c for c in table.columns if c not in ["Frequency", "Percentage"]][0]
        plot_df = table.sort_values("Frequency", ascending=True).tail(10)
        fig, ax = plt.subplots(figsize=(5.2, 3.5))
        bars = ax.barh(plot_df[category_col], plot_df["Frequency"])

        for bar, freq, pct in zip(bars, plot_df["Frequency"], plot_df["Percentage"]):
            ax.text(
                bar.get_width(),
                bar.get_y() + bar.get_height() / 2,
                f" {int(freq)} ({pct:.1f}%)",
                va="center",
                fontsize=7,
            )

        ax.set_title(label, fontsize=10, fontweight="bold")
        ax.set_xlabel("Frequency and percentage", fontsize=8)
        ax.tick_params(axis="both", labelsize=7)
        fig.tight_layout()

        outpath = output_dir / f"priority_{safe_filename(label)}.png"
        fig.savefig(outpath, dpi=240)
        plt.close(fig)
        saved_paths.append(outpath)

    return saved_paths


def render_priority_distribution_cards(priority_tables, detected_cols):
    """Display priority distribution tables in a polished two-column layout."""
    labels = list(priority_tables.keys())
    for i in range(0, len(labels), 2):
        cols = st.columns(2)
        for j, label in enumerate(labels[i:i + 2]):
            with cols[j]:
                table = priority_tables.get(label, pd.DataFrame())
                detected = detected_cols.get(label)
                st.markdown(f"#### {label}")
                if detected:
                    st.caption(f"Detected column: `{detected}`")
                if table is None or table.empty:
                    st.warning(f"No matching column found for {label}.")
                else:
                    st.dataframe(table, use_container_width=True, hide_index=True)


def render_clean_categorical_tables(df, categorical_cols, max_tables=20):
    """Display categorical distributions one variable at a time without a repeated Variable column."""
    shown = 0
    for col in categorical_cols:
        if shown >= max_tables:
            break
        table = variable_distribution_table(df, col)
        if table.empty:
            continue
        st.markdown(f"#### Distribution of Survey Participants by {pretty_label(col)}")
        st.dataframe(table, use_container_width=True, hide_index=True)
        shown += 1


# -------------------------------------------------------------------
# Descriptive analysis functions
# -------------------------------------------------------------------

def quantitative_summary(df, numeric_cols):
    """Produce average and descriptive statistics for all quantitative variables."""
    if not numeric_cols:
        return pd.DataFrame()

    summary = df[numeric_cols].describe().T.reset_index().rename(columns={
        "index": "Variable",
        "count": "Valid responses",
        "mean": "Mean / Average",
        "std": "Standard deviation",
        "min": "Minimum",
        "25%": "25th percentile",
        "50%": "Median",
        "75%": "75th percentile",
        "max": "Maximum"
    })

    numeric_summary_cols = [
        "Mean / Average", "Standard deviation", "Minimum",
        "25th percentile", "Median", "75th percentile", "Maximum"
    ]

    for col in numeric_summary_cols:
        if col in summary.columns:
            summary[col] = summary[col].round(2)

    return summary


def categorical_frequency_summary(df, categorical_cols):
    """Create frequency and percentage distribution for categorical variables."""
    all_results = []

    for col in categorical_cols:
        counts = df[col].fillna("Missing").astype(str).value_counts(dropna=False)
        total = counts.sum()

        for category, count in counts.items():
            all_results.append({
                "Variable": col,
                "Category": category,
                "Frequency": int(count),
                "Percentage": round((count / total) * 100, 2) if total > 0 else 0
            })

    return pd.DataFrame(all_results)


def likert_agree_summary(df, likert_cols):
    """Calculate percentage Agree or Strongly Agree for Likert variables coded 1-5."""
    results = []

    for col in likert_cols:
        series = pd.to_numeric(df[col], errors="coerce").dropna()
        total = len(series)

        if total == 0:
            continue

        agree_count = series[series >= 4].count()

        results.append({
            "Variable": col,
            "Valid responses": total,
            "Agree or Strongly Agree": agree_count,
            "Percentage Agree or Above": round((agree_count / total) * 100, 2)
        })

    return pd.DataFrame(results)


def yes_no_summary(df, yes_no_cols):
    """Calculate percentage Yes for Yes/No variables."""
    results = []

    for col in yes_no_cols:
        series = df[col].dropna().astype(str).str.lower().str.strip()
        total = len(series)

        if total == 0:
            continue

        yes_count = series.isin(["yes", "1", "true"]).sum()

        results.append({
            "Variable": col,
            "Valid responses": total,
            "Yes responses": int(yes_count),
            "Percentage Yes": round((yes_count / total) * 100, 2)
        })

    return pd.DataFrame(results)


def quantitative_summary_by_group(df, numeric_cols, group_col):
    """Mean/average quantitative analysis by gender or selected categorical group."""
    if group_col is None or group_col not in df.columns or not numeric_cols:
        return pd.DataFrame()

    results = []

    for group, gdf in df.groupby(group_col):
        for col in numeric_cols:
            series = pd.to_numeric(gdf[col], errors="coerce").dropna()
            if series.empty:
                continue

            results.append({
                "Group variable": group_col,
                "Group": group,
                "Variable": col,
                "Valid responses": len(series),
                "Mean / Average": round(series.mean(), 2),
                "Median": round(series.median(), 2),
                "Minimum": round(series.min(), 2),
                "Maximum": round(series.max(), 2),
                "Standard deviation": round(series.std(), 2) if len(series) > 1 else 0
            })

    return pd.DataFrame(results)


def likert_summary_by_group(df, likert_cols, group_col):
    """Likert Agree or Strongly Agree analysis by gender or selected categorical group."""
    if group_col is None or group_col not in df.columns or not likert_cols:
        return pd.DataFrame()

    results = []

    for group, gdf in df.groupby(group_col):
        for col in likert_cols:
            series = pd.to_numeric(gdf[col], errors="coerce").dropna()
            total = len(series)

            if total == 0:
                continue

            agree_count = series[series >= 4].count()

            results.append({
                "Group variable": group_col,
                "Group": group,
                "Variable": col,
                "Valid responses": total,
                "Agree or Strongly Agree": int(agree_count),
                "Percentage Agree or Above": round((agree_count / total) * 100, 2)
            })

    return pd.DataFrame(results)


def yes_no_summary_by_group(df, yes_no_cols, group_col):
    """Yes response percentage by gender or selected categorical group."""
    if group_col is None or group_col not in df.columns or not yes_no_cols:
        return pd.DataFrame()

    results = []

    for group, gdf in df.groupby(group_col):
        for col in yes_no_cols:
            series = gdf[col].dropna().astype(str).str.lower().str.strip()
            total = len(series)

            if total == 0:
                continue

            yes_count = series.isin(["yes", "1", "true"]).sum()

            results.append({
                "Group variable": group_col,
                "Group": group,
                "Variable": col,
                "Valid responses": total,
                "Yes responses": int(yes_count),
                "Percentage Yes": round((yes_count / total) * 100, 2)
            })

    return pd.DataFrame(results)


def categorical_frequency_by_group(df, categorical_cols, group_col):
    """Frequency distribution of categorical variables by gender or selected group."""
    if group_col is None or group_col not in df.columns or not categorical_cols:
        return pd.DataFrame()

    all_results = []

    for group, gdf in df.groupby(group_col):
        for col in categorical_cols:
            if col == group_col:
                continue

            counts = gdf[col].fillna("Missing").astype(str).value_counts(dropna=False)
            total = counts.sum()

            for category, count in counts.items():
                all_results.append({
                    "Group variable": group_col,
                    "Group": group,
                    "Variable": col,
                    "Category": category,
                    "Frequency": int(count),
                    "Percentage": round((count / total) * 100, 2) if total > 0 else 0
                })

    return pd.DataFrame(all_results)



def format_p_value(p_value):
    """Format p-values for reader-friendly tables."""
    if pd.isna(p_value):
        return "N/A"
    if p_value < 0.001:
        return "<0.001"
    return round(float(p_value), 4)


def significance_label(p_value):
    """Add a simple interpretation of statistical significance."""
    if pd.isna(p_value):
        return "Not tested"
    if p_value < 0.001:
        return "Highly significant"
    if p_value < 0.01:
        return "Very significant"
    if p_value < 0.05:
        return "Significant"
    return "Not significant"


def gender_quantitative_comparison(df_gender, numeric_cols, gender_group_col):
    """
    Variable-by-variable quantitative comparison table:
    Variable/Indicator | Men average | Women average | Difference | p-value.
    """
    if gender_group_col is None or df_gender.empty or not numeric_cols:
        return pd.DataFrame()

    results = []
    for col in numeric_cols:
        men = pd.to_numeric(
            df_gender.loc[df_gender[gender_group_col] == "Men", col], errors="coerce"
        ).dropna()
        women = pd.to_numeric(
            df_gender.loc[df_gender[gender_group_col] == "Women", col], errors="coerce"
        ).dropna()

        if men.empty and women.empty:
            continue

        men_mean = men.mean() if len(men) else pd.NA
        women_mean = women.mean() if len(women) else pd.NA
        difference = men_mean - women_mean if pd.notna(men_mean) and pd.notna(women_mean) else pd.NA

        p_value = pd.NA
        if stats is not None and len(men) >= 2 and len(women) >= 2:
            try:
                _, p_value = stats.ttest_ind(men, women, equal_var=False, nan_policy="omit")
            except Exception:
                p_value = pd.NA

        results.append({
            "Variable/Indicator": col,
            "Men (average/mean)": round(men_mean, 2) if pd.notna(men_mean) else "N/A",
            "Women (average/mean)": round(women_mean, 2) if pd.notna(women_mean) else "N/A",
            "Difference (Men - Women)": round(difference, 2) if pd.notna(difference) else "N/A",
            "Significance Level (p-value)": format_p_value(p_value),
            "Interpretation": significance_label(p_value),
            "Men n": int(len(men)),
            "Women n": int(len(women)),
        })

    return pd.DataFrame(results)


def quantitative_by_category_wide(df, numeric_cols, group_col):
    """
    For categorical variables with more than two categories, present the average of each
    quantitative variable by category and comparison statistics where appropriate.
    """
    if group_col is None or group_col not in df.columns or not numeric_cols:
        return pd.DataFrame()

    clean = df[[group_col] + numeric_cols].copy()
    clean[group_col] = clean[group_col].fillna("Missing").astype(str)
    categories = clean[group_col].dropna().unique().tolist()
    if len(categories) < 2:
        return pd.DataFrame()

    results = []
    for col in numeric_cols:
        row = {"Variable/Indicator": col, "Comparison variable": group_col}
        groups_for_test = []

        for category in sorted(categories):
            values = pd.to_numeric(clean.loc[clean[group_col] == category, col], errors="coerce").dropna()
            row[f"{category} mean"] = round(values.mean(), 2) if len(values) else "N/A"
            row[f"{category} n"] = int(len(values))
            if len(values) >= 2:
                groups_for_test.append(values)

        p_value = pd.NA
        if stats is not None and len(groups_for_test) >= 2:
            try:
                if len(groups_for_test) == 2:
                    _, p_value = stats.ttest_ind(groups_for_test[0], groups_for_test[1], equal_var=False, nan_policy="omit")
                else:
                    _, p_value = stats.f_oneway(*groups_for_test)
            except Exception:
                p_value = pd.NA

        row["Comparison statistic"] = "Welch t-test" if len(groups_for_test) == 2 else "One-way ANOVA"
        row["Significance Level (p-value)"] = format_p_value(p_value)
        row["Interpretation"] = significance_label(p_value)
        results.append(row)

    return pd.DataFrame(results)


def categorical_distribution_wide(df, group_col):
    """Frequency and percentage distribution for one categorical variable without repeating variable names."""
    return variable_distribution_table(df, group_col)


def top_summary_insights(quant_summary, likert_summary, yes_no_results, gender_comparison):
    """Generate short, reader-friendly insight bullets for the top of the dashboard."""
    insights = []
    if not quant_summary.empty:
        insights.append(f"Quantitative variables analyzed: {quant_summary['Variable'].nunique():,}.")
    if not likert_summary.empty:
        top = likert_summary.sort_values("Percentage Agree or Above", ascending=False).iloc[0]
        insights.append(
            f"Highest agreement barrier: {top['Variable']} ({top['Percentage Agree or Above']}% agree/strongly agree)."
        )
    if not yes_no_results.empty:
        top_yes = yes_no_results.sort_values("Percentage Yes", ascending=False).iloc[0]
        insights.append(f"Highest Yes response: {top_yes['Variable']} ({top_yes['Percentage Yes']}% Yes).")
    if not gender_comparison.empty and "Interpretation" in gender_comparison.columns:
        sig = gender_comparison[gender_comparison["Interpretation"].isin(["Significant", "Very significant", "Highly significant"])]
        insights.append(f"Gender comparison: {len(sig):,} quantitative variables show statistically significant Men/Women differences.")
    return insights


# -------------------------------------------------------------------
# Research question and cost-of-stalling analysis functions
# -------------------------------------------------------------------

def detect_first_column(df, candidates):
    """Find the first matching column from a list of possible aliases."""
    return find_column_by_candidates(df, candidates)


def contains_any(series, keywords):
    """Case-insensitive keyword match for categorical text."""
    s = series.fillna("").astype(str).str.lower()
    pattern = "|".join([re.escape(k.lower()) for k in keywords])
    return s.str.contains(pattern, regex=True, na=False)


def build_analysis_request_mapping(df, available_tables):
    """Map the requested analysis questions to current dashboard outputs."""
    checks = [
        ("Data readiness", "Number of observations, rows, columns, duplicates, and missing values", "key_statistics"),
        ("Participant profile", "Gender, year arrival, country origin, immigration status, marital status, education, employment status", "priority_distributions"),
        ("Quantitative descriptives", "Mean, median, minimum, maximum, standard deviation for all quantitative variables", "quantitative_summary"),
        ("Categorical distributions", "Frequency and percentage tables without repeated Variable column", "categorical_frequency"),
        ("Gender analysis", "Men/Women quantitative comparison with difference and p-value", "gender_quantitative_comparison"),
        ("Likert barriers", "Percent Agree or Strongly Agree for barrier variables", "likert_agree_summary"),
        ("Yes/No indicators", "Percent Yes for binary variables", "yes_no_summary"),
        ("SEM model", "Reliability, correlations, path coefficients, mediation, and SEM fit", "sem_estimates"),
        ("Cost of stalling", "Financial and social cost of stalled employment transition", "cost_of_stalling"),
        ("Research question alignment", "Mapping of requested questions to outputs", "research_question_mapping"),
    ]
    rows = []
    for area, requirement, table_key in checks:
        present = table_key in available_tables and available_tables.get(table_key) is not None and not available_tables.get(table_key, pd.DataFrame()).empty
        rows.append({
            "Analysis area": area,
            "Requirement from analysis request": requirement,
            "Provided by revised script": "Yes" if present else "Partial / depends on available variables",
            "Output location": table_key,
        })
    return pd.DataFrame(rows)


def identify_employment_stalling(df):
    """Create a flexible stalled-employment flag using available employment/job-search variables."""
    result = df.copy()
    employment_col = detect_first_column(result, [
        "employment_status", "employment_stat_us", "employment_stat", "current_employment_status",
        "work_status", "employment"
    ])
    months_col = detect_first_column(result, [
        "months_job_search", "job_search_months", "months_looking_for_work", "months_unemployed",
        "time_to_employment_months", "months_without_professional_job"
    ])
    underemployment_col = detect_first_column(result, [
        "underemployed", "underemployment", "working_below_skill", "job_below_qualification",
        "survival_job", "below_skill_level"
    ])

    stalled = pd.Series(False, index=result.index)
    evidence = []

    if employment_col:
        stalled_keywords = [
            "unemployed", "not employed", "looking", "job seeking", "seeking", "underemployed",
            "survival", "temporary", "part-time", "part time", "casual", "contract", "precarious"
        ]
        stalled = stalled | contains_any(result[employment_col], stalled_keywords)
        evidence.append(f"employment status column: {employment_col}")

    if months_col:
        months = pd.to_numeric(result[months_col], errors="coerce")
        stalled = stalled | (months >= 6)
        evidence.append(f"job-search duration >= 6 months: {months_col}")

    if underemployment_col:
        yes_values = {"yes", "1", "true", "underemployed", "below qualification", "below skill", "survival job"}
        stalled = stalled | result[underemployment_col].fillna("").astype(str).str.lower().str.strip().isin(yes_values)
        evidence.append(f"underemployment column: {underemployment_col}")

    result["stalled_employment_transition"] = stalled.map({True: "Stalled", False: "Not stalled / not flagged"})
    return result, evidence


def cost_of_stalling_analysis(df):
    """Estimate cost of stalled employment transition using all available financial/social-cost variables."""
    work_df, evidence = identify_employment_stalling(df)
    flag_col = "stalled_employment_transition"

    income_col = detect_first_column(work_df, [
        "annual_income_numeric", "annual_income", "monthly_income", "income", "current_income", "household_income", "personal_income",
        "employment_income", "income_monthly"
    ])
    expected_income_col = detect_first_column(work_df, [
        "expected_income", "expected_monthly_income", "target_income", "previous_income",
        "pre_arrival_income", "income_before_arrival", "professional_income_before_arrival"
    ])
    expense_col = detect_first_column(work_df, [
        "monthly_expenses", "expenses", "living_expenses", "household_expenses", "rent_burden",
        "monthly_cost", "cost_of_living"
    ])
    debt_col = detect_first_column(work_df, ["debt", "monthly_debt", "loan", "financial_debt"])
    stress_col = detect_first_column(work_df, ["financial_stress", "stress", "burnout", "mental_stress", "economic_stress"])

    numeric_cost_cols = [c for c in [income_col, expected_income_col, expense_col, debt_col, stress_col] if c]
    for c in numeric_cost_cols:
        work_df[c] = pd.to_numeric(work_df[c], errors="coerce")

    if income_col and expense_col:
        work_df["monthly_surplus_deficit"] = work_df[income_col] - work_df[expense_col]
        numeric_cost_cols.append("monthly_surplus_deficit")

    if income_col and expected_income_col:
        work_df["estimated_monthly_opportunity_cost"] = work_df[expected_income_col] - work_df[income_col]
        numeric_cost_cols.append("estimated_monthly_opportunity_cost")

    summary_rows = []
    for group, gdf in work_df.groupby(flag_col):
        row = {"Employment transition status": group, "Respondents": len(gdf)}
        for col in numeric_cost_cols:
            vals = pd.to_numeric(gdf[col], errors="coerce").dropna()
            row[f"{pretty_label(col)} mean"] = round(vals.mean(), 2) if len(vals) else "N/A"
            row[f"{pretty_label(col)} median"] = round(vals.median(), 2) if len(vals) else "N/A"
        summary_rows.append(row)

    summary = pd.DataFrame(summary_rows)

    comparison_rows = []
    for col in numeric_cost_cols:
        stalled_vals = pd.to_numeric(work_df.loc[work_df[flag_col] == "Stalled", col], errors="coerce").dropna()
        not_vals = pd.to_numeric(work_df.loc[work_df[flag_col] == "Not stalled / not flagged", col], errors="coerce").dropna()
        p_value = pd.NA
        if stats is not None and len(stalled_vals) >= 2 and len(not_vals) >= 2:
            try:
                _, p_value = stats.ttest_ind(stalled_vals, not_vals, equal_var=False, nan_policy="omit")
            except Exception:
                p_value = pd.NA
        comparison_rows.append({
            "Cost indicator": pretty_label(col),
            "Stalled mean": round(stalled_vals.mean(), 2) if len(stalled_vals) else "N/A",
            "Not stalled mean": round(not_vals.mean(), 2) if len(not_vals) else "N/A",
            "Difference": round(stalled_vals.mean() - not_vals.mean(), 2) if len(stalled_vals) and len(not_vals) else "N/A",
            "p-value": format_p_value(p_value),
            "Interpretation": significance_label(p_value),
            "Stalled n": int(len(stalled_vals)),
            "Not stalled n": int(len(not_vals)),
        })
    comparison = pd.DataFrame(comparison_rows)

    flag_distribution = variable_distribution_table(work_df, flag_col)
    metadata = pd.DataFrame({
        "Item": ["Stalling definition evidence", "Income variable", "Expected / previous income variable", "Expense variable", "Debt variable", "Stress variable"],
        "Detected value": ["; ".join(evidence) if evidence else "No direct stalling variable detected; flag defaults to employment/job-search proxies only where available", income_col or "Not detected", expected_income_col or "Not detected", expense_col or "Not detected", debt_col or "Not detected", stress_col or "Not detected"]
    })
    return work_df, summary, comparison, flag_distribution, metadata


def barrier_construct_alignment(df):
    """Create an analysis table for common stalled-employment barriers using available columns."""
    barrier_specs = {
        "Credential recognition barrier": ["credential_recognition", "cred_recognition", "cred_req_difficulty", "credential_difficulty", "foreign_credential"],
        "Canadian experience barrier": ["canadian_experience", "lack_canadian_experience", "canadian_exp_barrier", "canadian_work_experience"],
        "Discrimination in hiring": ["disc_hiring", "discrimination", "hiring_discrimination", "workplace_discrimination"],
        "Language barrier": ["language_barrier", "english_barrier", "french_barrier", "language_difficulty"],
        "Professional network barrier": ["professional_contacts", "networking", "social_capital", "professional_network"],
        "Financial pressure": ["financial_stress", "financial_pressure", "rent_burden", "cost_of_living"],
        "Mental wellbeing pressure": ["burnout", "stress", "mental_health", "wellbeing"],
    }
    rows = []
    for barrier, aliases in barrier_specs.items():
        col = detect_first_column(df, aliases)
        if col is None:
            rows.append({"Barrier / construct": barrier, "Detected column": "Not detected", "Analysis produced": "No", "Key statistic": "N/A"})
            continue
        ser_num = pd.to_numeric(df[col], errors="coerce")
        if ser_num.notna().sum() > 0 and ser_num.dropna().between(1, 5).all():
            valid = ser_num.dropna()
            agree = (valid >= 4).sum()
            stat_text = f"{round((agree / len(valid)) * 100, 2)}% agree/strongly agree"
        else:
            vals = df[col].fillna("Missing").astype(str).str.lower().str.strip()
            yes = vals.isin(["yes", "1", "true", "agree", "strongly agree"]).sum()
            stat_text = f"{round((yes / len(vals)) * 100, 2)}% yes/positive" if len(vals) else "N/A"
        rows.append({"Barrier / construct": barrier, "Detected column": col, "Analysis produced": "Yes", "Key statistic": stat_text})
    return pd.DataFrame(rows)

# -------------------------------------------------------------------
# Visualization functions
# -------------------------------------------------------------------

def save_likert_bar_chart(summary_df, outpath):
    if summary_df.empty:
        return

    fig, ax = plt.subplots(figsize=(9, 5))
    plot_df = summary_df.sort_values("Percentage Agree or Above", ascending=True).tail(15)

    ax.barh(plot_df["Variable"], plot_df["Percentage Agree or Above"])
    ax.set_xlabel("Percentage Agree or Strongly Agree")
    ax.set_title("Top Likert-Scale Barriers")
    fig.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)


def save_grouped_likert_chart(summary_df, outpath, title="Gender-Disaggregated Likert Analysis"):
    if summary_df.empty:
        return

    plot_df = summary_df.copy()

    # Keep top variables by average agreement for readability
    top_vars = (
        plot_df.groupby("Variable")["Percentage Agree or Above"]
        .mean()
        .sort_values(ascending=False)
        .head(10)
        .index
    )

    plot_df = plot_df[plot_df["Variable"].isin(top_vars)]

    fig, ax = plt.subplots(figsize=(10, 5))
    pivot = plot_df.pivot_table(
        index="Variable",
        columns="Group",
        values="Percentage Agree or Above",
        aggfunc="mean"
    ).fillna(0)

    pivot.plot(kind="barh", ax=ax)
    ax.set_xlabel("Percentage Agree or Strongly Agree")
    ax.set_title(title)
    ax.legend(title="Group", loc="best")
    fig.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)


def save_numeric_histograms(df, numeric_cols, output_dir):
    """Save histograms with both absolute counts and percentage labels."""
    saved_paths = []

    for col in numeric_cols[:12]:
        series = pd.to_numeric(df[col], errors="coerce").dropna()

        if series.empty:
            continue

        fig, ax = plt.subplots(figsize=(4.8, 3.4))
        counts, bins, patches = ax.hist(series, bins=12)
        total = counts.sum()

        for count, patch in zip(counts, patches):
            if count <= 0 or total <= 0:
                continue
            pct = (count / total) * 100
            x = patch.get_x() + patch.get_width() / 2
            y = patch.get_height()
            ax.text(x, y, f"{int(count)}\n({pct:.1f}%)", ha="center", va="bottom", fontsize=6)

        ax.set_title(f"{col}", fontsize=9, fontweight="bold")
        ax.set_xlabel(col, fontsize=8)
        ax.set_ylabel("Frequency", fontsize=8)
        ax.tick_params(axis="both", labelsize=7)
        fig.tight_layout()

        outpath = output_dir / f"histogram_{safe_filename(col)}.png"
        fig.savefig(outpath, dpi=240)
        plt.close(fig)
        saved_paths.append(outpath)

    return saved_paths


def save_yes_no_pie_charts(df, yes_no_cols, output_dir):
    saved_paths = []

    for col in yes_no_cols[:8]:
        series = df[col].dropna().astype(str).str.lower().str.strip()
        yes_count = series.isin(["yes", "1", "true"]).sum()
        no_count = series.isin(["no", "0", "false"]).sum()

        if yes_count + no_count == 0:
            continue

        fig, ax = plt.subplots(figsize=(4.2, 3.2))
        ax.pie([yes_count, no_count], labels=["Yes", "No"], autopct="%1.1f%%", textprops={"fontsize": 8})
        ax.set_title(f"{col}", fontsize=9)

        outpath = output_dir / f"pie_{safe_filename(col)}.png"
        fig.savefig(outpath, dpi=220)
        plt.close(fig)
        saved_paths.append(outpath)

    return saved_paths


def save_categorical_bar_charts(df, categorical_cols, output_dir):
    """Save categorical bar charts with frequency and percentage labels."""
    saved_paths = []

    for col in categorical_cols[:12]:
        counts = df[col].fillna("Missing").astype(str).value_counts().head(10)

        if counts.empty:
            continue

        total = counts.sum()
        plot_counts = counts.sort_values()
        fig, ax = plt.subplots(figsize=(5.0, 3.4))
        bars = ax.barh(plot_counts.index, plot_counts.values)

        for bar, value in zip(bars, plot_counts.values):
            pct = (value / total) * 100 if total > 0 else 0
            ax.text(
                bar.get_width(),
                bar.get_y() + bar.get_height() / 2,
                f" {int(value)} ({pct:.1f}%)",
                va="center",
                fontsize=7,
            )

        ax.set_title(f"{col}", fontsize=9, fontweight="bold")
        ax.set_xlabel("Frequency and percentage", fontsize=8)
        ax.tick_params(axis="both", labelsize=7)
        fig.tight_layout()

        outpath = output_dir / f"categorical_{safe_filename(col)}.png"
        fig.savefig(outpath, dpi=240)
        plt.close(fig)
        saved_paths.append(outpath)

    return saved_paths


def show_figures_grid(image_paths, columns_per_row=4):
    """Display smaller figures with at least four figures per row."""
    if not image_paths:
        st.info("No figures available for this section.")
        return

    for i in range(0, len(image_paths), columns_per_row):
        cols = st.columns(columns_per_row)
        for j, path in enumerate(image_paths[i:i + columns_per_row]):
            with cols[j]:
                st.image(str(path), use_container_width=True)


def save_corr_heatmap(corr: pd.DataFrame, outpath: Path):
    fig, ax = plt.subplots(figsize=(9, 6))
    sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm", center=0, ax=ax)
    ax.set_title("Correlation Matrix")
    fig.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)


def save_path_diagram(paths: pd.DataFrame, outpath: Path):
    fig, ax = plt.subplots(figsize=(10, 7))
    g = nx.DiGraph()

    for _, row in paths.iterrows():
        g.add_edge(row["From"], row["To"], weight=row["Beta/std r"])

    if not g.nodes:
        ax.text(0.5, 0.5, "No valid path coefficients available", ha="center", va="center")
    else:
        pos = nx.spring_layout(g, seed=42, k=1.3)
        nx.draw_networkx_nodes(g, pos, node_size=2200, node_color="#f2f6ff", edgecolors="#333333", ax=ax)
        nx.draw_networkx_labels(g, pos, font_size=10, ax=ax)
        nx.draw_networkx_edges(g, pos, arrows=True, arrowstyle="-|>", arrowsize=18, width=1.8, ax=ax)

        labels = {
            (r["From"], r["To"]): f"{r['Beta/std r']:.2f}"
            for _, r in paths.iterrows()
        }
        nx.draw_networkx_edge_labels(g, pos, edge_labels=labels, font_size=9, ax=ax)

    ax.set_axis_off()
    fig.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)


# -------------------------------------------------------------------
# Word report
# -------------------------------------------------------------------

def create_word_report(tables, figures, outpath: Path):
    doc = Document()

    doc.add_heading("EMERGE+ Newcomer SEM Analysis Report", 0)

    doc.add_paragraph(
        "This automated report summarizes data readiness, descriptive statistics, "
        "categorical frequency distributions, gender-disaggregated findings, construct reliability, "
        "correlations, path coefficients, mediation tests, and SEM model fit where available."
    )

    doc.add_heading("1. Variable Mapping", level=1)
    add_docx_table(doc, tables["variable_mapping"].head(40))

    doc.add_heading("2. Descriptive Analysis", level=1)

    doc.add_heading("2.0 Key Dataset Statistics", level=2)
    add_docx_table(doc, tables.get("key_statistics", pd.DataFrame()))

    doc.add_heading("2.0b Missing Values Summary", level=2)
    add_docx_table(doc, tables.get("missing_values_summary", pd.DataFrame()).head(40))

    doc.add_heading("2.0c Priority Demographic and Background Distributions", level=2)
    add_docx_table(doc, tables.get("priority_distributions", pd.DataFrame()).head(80))

    doc.add_heading("2.1 Quantitative Variables: Averages and Distribution", level=2)
    add_docx_table(doc, tables["quantitative_summary"].head(40))

    doc.add_heading("2.2 Participant Profile Distributions", level=2)
    for key, value in tables.get("_key_distribution_tables", {}).items():
        doc.add_heading(f"Distribution of Survey Participants by {key}", level=3)
        add_docx_table(doc, value["table"])

    doc.add_heading("2.3 Full Categorical Frequency Distribution", level=2)
    doc.add_paragraph("The dashboard presents category-specific tables to avoid repeating the variable name in every row.")
    add_docx_table(doc, tables["categorical_frequency"].head(60))

    doc.add_heading("2.4 Likert-Scale Agree or Above Analysis", level=2)
    add_docx_table(doc, tables["likert_agree_summary"].head(40))

    if figures.get("likert_bar") and figures["likert_bar"].exists():
        doc.add_picture(str(figures["likert_bar"]), width=Inches(6.5))

    doc.add_heading("2.4 Yes/No Variables", level=2)
    add_docx_table(doc, tables["yes_no_summary"].head(40))

    doc.add_heading("3. Gender-Disaggregated Analysis: Men and Women Only", level=1)
    doc.add_paragraph(
        "This section includes only respondents coded as Men or Women. Other gender responses "
        "are excluded from this specific disaggregation to meet the requested reporting format."
    )

    doc.add_heading("3.1 Variable-by-Variable Quantitative Comparison by Gender", level=2)
    add_docx_table(doc, tables["gender_quantitative_comparison"].head(80))

    doc.add_heading("3.1b Detailed Quantitative Averages by Gender", level=2)
    add_docx_table(doc, tables["gender_quantitative"].head(60))

    doc.add_heading("3.2 Likert Agreement by Gender", level=2)
    add_docx_table(doc, tables["gender_likert"].head(60))

    if figures.get("gender_likert") and figures["gender_likert"].exists():
        doc.add_picture(str(figures["gender_likert"]), width=Inches(6.5))

    doc.add_heading("3.3 Yes/No Responses by Gender", level=2)
    add_docx_table(doc, tables["gender_yes_no"].head(60))

    doc.add_heading("3.4 Categorical Frequency by Gender", level=2)
    add_docx_table(doc, tables["gender_categorical"].head(60))

    doc.add_heading("4. Reliability Summary", level=1)
    add_docx_table(doc, tables["reliability"])

    doc.add_heading("5. Correlation Heatmap", level=1)
    if figures.get("correlation") and figures["correlation"].exists():
        doc.add_picture(str(figures["correlation"]), width=Inches(6.5))

    doc.add_heading("6. SEM Path Diagram", level=1)
    if figures.get("paths") and figures["paths"].exists():
        doc.add_picture(str(figures["paths"]), width=Inches(6.5))

    doc.add_heading("7. Path Coefficients", level=1)
    add_docx_table(doc, tables["paths"].head(20))

    doc.add_heading("8. Mediation Analysis", level=1)
    add_docx_table(doc, tables["mediation"].head(20))

    doc.add_heading("9. Interpretation Notes", level=1)
    doc.add_paragraph(
        "Likert-scale results show the percentage of respondents who agreed or strongly agreed. "
        "Yes/No results show the share of respondents who answered Yes. Quantitative results show averages "
        "and other descriptive statistics. Gender-disaggregated results compare Men and Women only."
    )

    doc.save(outpath)
    return outpath


# -------------------------------------------------------------------
# Streamlit app
# -------------------------------------------------------------------

uploaded = st.file_uploader("Upload EMERGE newcomer dataset", type=["csv", "xlsx", "xls"])

if not uploaded:
    st.info(
        "Upload your dataset to begin. The app expects variables such as "
        "cred_req_difficulty, disc_hiring, rent_burden, professional_contacts, "
        "burnout, online_job_confidence, and employment_confidence."
    )
    st.stop()

raw = load_data(uploaded)
df = normalize_columns(raw)
df = coerce_analysis_columns(df)
# Convert income bands such as 25_49k and 100k_plus into numeric annual income values for analysis.
df = add_income_numeric_columns(df)

# Variable mapping is calculated early but displayed below the key findings.
mapping = variable_mapping(df)

numeric_cols = detect_numeric_variables(df)
likert_cols = detect_likert_variables(df)
yes_no_cols = detect_yes_no_variables(df)
categorical_cols = detect_categorical_variables(df, numeric_cols=numeric_cols)

gender_col = find_gender_column(df)
df_gender, gender_group_col = standardize_gender_for_men_women(df, gender_col)

st.sidebar.header("Reader-Friendly Analysis Options")
st.sidebar.caption("Choose additional grouping variables if you want analysis beyond gender.")

available_group_vars = [
    col for col in categorical_cols
    if col != gender_col and df[col].nunique(dropna=True) <= 20
]

selected_group_vars = st.sidebar.multiselect(
    "Optional: choose other categorical variables for disaggregated analysis",
    options=available_group_vars,
    default=[]
)

# Original SEM workflow remains unchanged
scores = construct_scores(df)
desc = descriptive_stats(df)
rel, loadings = reliability_table(df)
corr = correlation_matrix(scores)
paths = path_coefficients(scores)
med = mediation_analysis(scores)
_, sem_estimates, sem_fit = semopy_model(scores)

# New descriptive summaries
quant_summary = quantitative_summary(df, numeric_cols)
categorical_freq = categorical_frequency_summary(df, categorical_cols)
key_distribution_tables = build_key_distribution_tables(df, gender_col=gender_col)
likert_summary = likert_agree_summary(df, likert_cols)
yes_no_results = yes_no_summary(df, yes_no_cols)

# Gender-disaggregated summaries: Men and Women only
gender_quant = quantitative_summary_by_group(df_gender, numeric_cols, gender_group_col)
gender_quant_comparison = gender_quantitative_comparison(df_gender, numeric_cols, gender_group_col)
gender_likert = likert_summary_by_group(df_gender, likert_cols, gender_group_col)
gender_yes_no = yes_no_summary_by_group(df_gender, yes_no_cols, gender_group_col)
gender_categorical = categorical_frequency_by_group(df_gender, categorical_cols, gender_group_col)

# Key dataset statistics and requested demographic/background distributions
key_stats_table = key_dataset_statistics(df)
missing_summary = missing_values_summary(df)
priority_tables, detected_priority_columns = priority_distribution_tables(df, gender_col=gender_col)
priority_distribution_summary = combine_priority_distributions(priority_tables)

# Research-question alignment and cost-of-stalling analysis
stalling_df, cost_summary, cost_comparison, stalling_distribution, cost_metadata = cost_of_stalling_analysis(df)
barrier_alignment = barrier_construct_alignment(df)

# Optional categorical disaggregation selected by client/user
optional_group_tables = {}
for group_col in selected_group_vars:
    optional_group_tables[f"{group_col}_quantitative_comparison"] = quantitative_by_category_wide(df, numeric_cols, group_col)
    optional_group_tables[f"{group_col}_distribution"] = categorical_distribution_wide(df, group_col)
    optional_group_tables[f"{group_col}_quantitative"] = quantitative_summary_by_group(df, numeric_cols, group_col)
    optional_group_tables[f"{group_col}_likert"] = likert_summary_by_group(df, likert_cols, group_col)
    optional_group_tables[f"{group_col}_yes_no"] = yes_no_summary_by_group(df, yes_no_cols, group_col)
    optional_group_tables[f"{group_col}_categorical"] = categorical_frequency_by_group(df, categorical_cols, group_col)

# Save figures
fig_corr = OUTPUT_ROOT / "figures" / "correlation_heatmap.png"
fig_paths = OUTPUT_ROOT / "figures" / "sem_path_diagram.png"
fig_likert = OUTPUT_ROOT / "figures" / "likert_bar_chart.png"
fig_gender_likert = OUTPUT_ROOT / "figures" / "gender_likert_bar_chart.png"

if not corr.empty:
    save_corr_heatmap(corr, fig_corr)

save_path_diagram(paths, fig_paths)
save_likert_bar_chart(likert_summary, fig_likert)
save_grouped_likert_chart(gender_likert, fig_gender_likert)

histogram_paths = save_numeric_histograms(df, numeric_cols, OUTPUT_ROOT / "figures")
pie_paths = save_yes_no_pie_charts(df, yes_no_cols, OUTPUT_ROOT / "figures")
categorical_chart_paths = save_categorical_bar_charts(df, categorical_cols, OUTPUT_ROOT / "figures")
priority_chart_paths = save_priority_distribution_charts(priority_tables, OUTPUT_ROOT / "figures")

# Save processed outputs
processed_path = OUTPUT_ROOT / "processed_data" / "cleaned_emerge_dataset.csv"
scores_path = OUTPUT_ROOT / "processed_data" / "construct_scores.csv"

df.to_csv(processed_path, index=False)
scores.to_csv(scores_path, index=False)

all_tables = {
    "variable_mapping": mapping,
    "descriptive_stats": desc,
    "quantitative_summary": quant_summary,
    "key_statistics": key_stats_table,
    "missing_values_summary": missing_summary,
    "priority_distributions": priority_distribution_summary,
    "cost_of_stalling": cost_summary,
    "cost_of_stalling_comparison": cost_comparison,
    "stalling_distribution": stalling_distribution,
    "cost_of_stalling_metadata": cost_metadata,
    "barrier_construct_alignment": barrier_alignment,
    "categorical_frequency": categorical_freq,
    "likert_agree_summary": likert_summary,
    "yes_no_summary": yes_no_results,
    "gender_quantitative": gender_quant,
    "gender_quantitative_comparison": gender_quant_comparison,
    "gender_likert": gender_likert,
    "gender_yes_no": gender_yes_no,
    "gender_categorical": gender_categorical,
    "reliability": rel,
    "factor_loadings": loadings,
    "correlation_matrix": corr.reset_index(),
    "paths": paths,
    "mediation": med,
    "sem_estimates": sem_estimates,
    "sem_fit": sem_fit,
}

research_question_mapping = build_analysis_request_mapping(df, all_tables)
all_tables["research_question_mapping"] = research_question_mapping

# Add clean participant distribution tables to the Excel workbook.
# Each table has only: category name, Frequency, and Percentage.
for dist_title, dist_info in key_distribution_tables.items():
    sheet_name = f"distribution_{safe_filename(dist_title).lower()}"
    all_tables[sheet_name] = dist_info["table"]

# Add optional selected group analysis into Excel workbook
all_tables.update(optional_group_tables)

excel_path = OUTPUT_ROOT / "tables" / "emerge_sem_results.xlsx"
write_excel(all_tables, excel_path)

report_path = OUTPUT_ROOT / "model_outputs" / "emerge_sem_report.docx"
report_tables = dict(all_tables)
report_tables["_key_distribution_tables"] = key_distribution_tables
create_word_report(
    report_tables,
    {
        "correlation": fig_corr,
        "paths": fig_paths,
        "likert_bar": fig_likert,
        "gender_likert": fig_gender_likert,
    },
    report_path
)

# -------------------------------------------------------------------
# Display results
# -------------------------------------------------------------------

st.markdown("""
<div class="hero-card">
<h2>EMERGE+ Newcomer SEM Dashboard</h2>
<p>Key analysis results, data readiness, demographic distribution, descriptive statistics, gender analysis, and SEM findings are presented in a reader-friendly dashboard.</p>
</div>
""", unsafe_allow_html=True)

st.markdown("## Analysis Results")
st.caption("Key findings are shown first so users can immediately understand the dataset before exploring detailed tabs.")

insights = top_summary_insights(quant_summary, likert_summary, yes_no_results, gender_quant_comparison)

kpi1, kpi2, kpi3, kpi4 = st.columns(4)
with kpi1:
    st.markdown(f'<div class="metric-card"><h3>Rows analyzed</h3><p>{df.shape[0]:,}</p></div>', unsafe_allow_html=True)
with kpi2:
    st.markdown(f'<div class="metric-card"><h3>Variables</h3><p>{df.shape[1]:,}</p></div>', unsafe_allow_html=True)
with kpi3:
    st.markdown(f'<div class="metric-card"><h3>Quantitative variables</h3><p>{len(numeric_cols):,}</p></div>', unsafe_allow_html=True)
with kpi4:
    st.markdown(f'<div class="metric-card"><h3>Categorical variables</h3><p>{len(categorical_cols):,}</p></div>', unsafe_allow_html=True)

missing_cells_home = int(df.isna().sum().sum())
missing_pct_home = round((missing_cells_home / (df.shape[0] * df.shape[1])) * 100, 2) if df.shape[0] * df.shape[1] else 0
st.markdown(
    f'<div class="stat-strip"><b>Data readiness:</b> {df.shape[0]:,} observations/rows, {df.shape[1]:,} columns, '
    f'{int(df.duplicated().sum()):,} duplicated rows, {missing_cells_home:,} missing values ({missing_pct_home}% of all cells).</div>',
    unsafe_allow_html=True
)

if insights:
    st.markdown('<div class="insight-box"><b>Summary insights</b><br>' + '<br>'.join([f'• {item}' for item in insights]) + '</div>', unsafe_allow_html=True)

st.markdown("### Key Demographic and Background Distributions")
st.caption("Requested distributions: gender, year of arrival, country of origin, immigration status, marital status, highest education, and employment status.")
render_priority_distribution_cards(priority_tables, detected_priority_columns)

st.markdown("### Key Quantitative Comparison by Gender")
if gender_quant_comparison.empty:
    st.info("Gender comparison table is not available. Check that the dataset has a valid gender/sex column coded for Men and Women.")
else:
    st.dataframe(gender_quant_comparison, use_container_width=True)

st.markdown("### Top Descriptive Findings")
c_top1, c_top2 = st.columns(2)
with c_top1:
    st.markdown("#### Quantitative Averages")
    st.dataframe(quant_summary[["Variable", "Valid responses", "Mean / Average", "Median", "Minimum", "Maximum"]].head(12), use_container_width=True)
with c_top2:
    st.markdown("#### Participant Profile Distribution")
    if key_distribution_tables:
        first_title, first_info = next(iter(key_distribution_tables.items()))
        st.caption(f"Distribution of survey participants by {first_title}")
        st.dataframe(first_info["table"], use_container_width=True, hide_index=True)
    else:
        st.info("No participant profile distribution variables detected.")

st.markdown("### Dashboard Details")

tabs = st.tabs([
    "Overview & Data Readiness",
    "Key Statistics",
    "Request Mapping",
    "Cost of Stalling",
    "Descriptives",
    "Categorical Frequencies",
    "Gender Analysis",
    "Optional Group Analysis",
    "Visuals",
    "Reliability",
    "Correlations",
    "Paths",
    "Mediation",
    "SEM fit"
])

with tabs[0]:
    st.subheader("Dataset Overview")
    col1, col2, col3 = st.columns(3)
    col1.metric("Rows", f"{df.shape[0]:,}")
    col2.metric("Columns", f"{df.shape[1]:,}")
    col3.metric("Duplicated rows", int(df.duplicated().sum()))
    st.dataframe(df.head(), use_container_width=True)

    st.subheader("Variable Mapping")
    st.dataframe(mapping, use_container_width=True)

with tabs[1]:
    st.subheader("Key Dataset Statistics")
    st.caption("This section summarizes the number of observations, rows, columns, missing values, and the main demographic/background distributions requested.")
    st.dataframe(key_stats_table, use_container_width=True, hide_index=True)

    st.subheader("Missing Values by Variable")
    st.dataframe(missing_summary, use_container_width=True, hide_index=True)

    st.subheader("Requested Demographic and Background Distributions")
    render_priority_distribution_cards(priority_tables, detected_priority_columns)

    st.subheader("Priority Distribution Charts")
    show_figures_grid(priority_chart_paths, columns_per_row=4)


with tabs[2]:
    st.subheader("Mapping of Requested Data Analysis Questions to Dashboard Outputs")
    st.caption("This table verifies whether the script produces each required analysis output. Items marked partial depend on whether the uploaded dataset contains the required variables.")
    st.dataframe(research_question_mapping, use_container_width=True, hide_index=True)

    st.subheader("Barrier / Construct Alignment")
    st.caption("This checks whether key stalled-employment barriers are detected in the uploaded dataset and whether summary statistics can be produced.")
    st.dataframe(barrier_alignment, use_container_width=True, hide_index=True)

with tabs[3]:
    st.subheader("Cost of Stalled Employment Transition")
    st.caption("This section estimates the cost of stalled employment transition using available employment status, job-search duration, income, expenses, debt, and stress variables.")

    st.markdown("#### Stalled Employment Transition Distribution")
    st.dataframe(stalling_distribution, use_container_width=True, hide_index=True)

    st.markdown("#### Variables Detected for Cost-of-Stalling Analysis")
    st.dataframe(cost_metadata, use_container_width=True, hide_index=True)

    st.markdown("#### Cost Summary by Employment Transition Status")
    st.dataframe(cost_summary, use_container_width=True, hide_index=True)

    st.markdown("#### Cost Comparison: Stalled vs Not Stalled")
    st.caption("Where numeric data are available, this table compares mean values and provides p-values using Welch t-tests.")
    st.dataframe(cost_comparison, use_container_width=True, hide_index=True)

with tabs[4]:
    st.subheader("Quantitative Variables: Average and Descriptive Statistics")
    st.caption("This table provides averages and descriptive statistics for all quantitative variables.")
    st.dataframe(quant_summary, use_container_width=True)

    st.subheader("Existing Descriptive Statistics")
    st.dataframe(desc, use_container_width=True)

with tabs[5]:
    st.subheader("Frequency Distribution for Categorical Variables")
    st.caption("Each categorical variable is shown in its own clean table. The repeated 'Variable' column has been removed from the dashboard display.")
    render_clean_categorical_tables(df, categorical_cols, max_tables=25)

    st.subheader("Likert-Scale Agree or Strongly Agree Analysis")
    st.dataframe(likert_summary, use_container_width=True)

    st.subheader("Yes/No Response Analysis")
    st.dataframe(yes_no_results, use_container_width=True)

with tabs[6]:
    st.subheader("Gender-Disaggregated Analysis: Men and Women Only")

    if gender_col is None:
        st.warning("No gender column was detected. Please ensure the dataset has a column named gender, sex, respondent_gender, or participant_gender.")
    elif df_gender.empty:
        st.warning("A gender column was detected, but no valid Men/Women records were found after recoding.")
    else:
        st.caption(f"Detected gender column: {gender_col}. Analysis below excludes third-option or other responses for this specific gender comparison.")

        st.markdown("#### Variable-by-Variable Quantitative Comparison by Gender")
        st.caption("This table follows the requested format: Variable/Indicator, Men average, Women average, Difference, and p-value.")
        st.dataframe(gender_quant_comparison, use_container_width=True)

        st.markdown("#### Detailed Quantitative Averages by Gender")
        st.dataframe(gender_quant, use_container_width=True)

        st.markdown("#### Likert Agreement by Gender")
        st.dataframe(gender_likert, use_container_width=True)

        if fig_gender_likert.exists():
            st.image(str(fig_gender_likert), use_container_width=True)

        st.markdown("#### Yes/No Responses by Gender")
        st.dataframe(gender_yes_no, use_container_width=True)

        st.markdown("#### Categorical Frequencies by Gender")
        st.dataframe(gender_categorical, use_container_width=True)

with tabs[7]:
    st.subheader("Optional Disaggregated Analysis by Selected Categorical Variables")
    st.caption("Use the sidebar to choose categorical variables such as region, education, age group, country of origin, or employment status.")

    if not selected_group_vars:
        st.info("No additional categorical variable selected. Use the sidebar to choose one or more variables.")
    else:
        for group_col in selected_group_vars:
            st.markdown(f"### Analysis by {group_col}")

            st.markdown("#### Quantitative Comparison by Category")
            st.caption("For categorical variables with two or more groups, this table presents mean values by category and includes p-values where appropriate.")
            st.dataframe(optional_group_tables.get(f"{group_col}_quantitative_comparison", pd.DataFrame()), use_container_width=True)

            st.markdown("#### Category Distribution")
            st.dataframe(optional_group_tables.get(f"{group_col}_distribution", pd.DataFrame()), use_container_width=True)

            st.markdown("#### Detailed Quantitative Averages")
            st.dataframe(optional_group_tables.get(f"{group_col}_quantitative", pd.DataFrame()), use_container_width=True)

            st.markdown("#### Likert Agreement")
            st.dataframe(optional_group_tables.get(f"{group_col}_likert", pd.DataFrame()), use_container_width=True)

            st.markdown("#### Yes/No Responses")
            st.dataframe(optional_group_tables.get(f"{group_col}_yes_no", pd.DataFrame()), use_container_width=True)

            st.markdown("#### Categorical Frequency")
            st.dataframe(optional_group_tables.get(f"{group_col}_categorical", pd.DataFrame()), use_container_width=True)

with tabs[8]:
    st.subheader("Reader-Friendly Visual Summary")
    st.caption("Figures are shown in a compact layout with four visuals per row.")

    st.markdown("#### Priority Demographic and Background Charts")
    st.caption("Compact charts for gender, year of arrival, country of origin, immigration status, marital status, highest education, and employment status where available.")
    show_figures_grid(priority_chart_paths, columns_per_row=4)

    st.markdown("#### Quantitative Histograms: Absolute Counts and Percentages")
    st.caption("Each histogram labels bars with both number of respondents and percentage share.")
    show_figures_grid(histogram_paths, columns_per_row=4)

    st.markdown("#### Categorical Bar Charts: Frequency and Percentage")
    show_figures_grid(categorical_chart_paths, columns_per_row=4)

    st.markdown("#### Yes/No Pie Charts")
    show_figures_grid(pie_paths, columns_per_row=4)

    st.markdown("#### Key Likert and Gender Visuals")
    main_visuals = [p for p in [fig_likert, fig_gender_likert, fig_corr, fig_paths] if p.exists()]
    show_figures_grid(main_visuals, columns_per_row=4)

with tabs[9]:
    st.dataframe(rel, use_container_width=True)
    st.dataframe(loadings, use_container_width=True)

with tabs[10]:
    if fig_corr.exists():
        st.image(str(fig_corr), use_container_width=True)
    st.dataframe(corr, use_container_width=True)

with tabs[11]:
    if fig_paths.exists():
        st.image(str(fig_paths), use_container_width=True)
    st.dataframe(paths, use_container_width=True)

with tabs[12]:
    st.dataframe(med, use_container_width=True)

with tabs[13]:
    st.dataframe(sem_estimates, use_container_width=True)
    st.dataframe(sem_fit, use_container_width=True)

# -------------------------------------------------------------------
# Download outputs
# -------------------------------------------------------------------

st.subheader("Download outputs")

c1, c2, c3, c4 = st.columns(4)

with open(excel_path, "rb") as f:
    c1.download_button(
        "Download Excel tables",
        f,
        file_name="emerge_sem_results.xlsx"
    )

with open(report_path, "rb") as f:
    c2.download_button(
        "Download Word report",
        f,
        file_name="emerge_sem_report.docx"
    )

with open(processed_path, "rb") as f:
    c3.download_button(
        "Download cleaned CSV",
        f,
        file_name="cleaned_emerge_dataset.csv"
    )

with open(scores_path, "rb") as f:
    c4.download_button(
        "Download construct scores",
        f,
        file_name="construct_scores.csv"
    )
